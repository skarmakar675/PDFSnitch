from __future__ import annotations

import base64
import html
import io
import logging
import os
import re
import shutil
import subprocess
import tempfile
import uuid
import zipfile
from contextlib import contextmanager
from importlib import metadata as importlib_metadata
from importlib import util as importlib_util
from pathlib import Path
from typing import Any, Iterable

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response
from PIL import Image, ImageEnhance, ImageOps, UnidentifiedImageError
from pypdf import PdfReader, PdfWriter
from reportlab.lib import colors
from reportlab.lib.colors import Color
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfgen import canvas
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from starlette.concurrency import run_in_threadpool

logger = logging.getLogger("pdfsnitch")
_PADDLE_OCR_ENGINE: Any | None = None
_PADDLE_OCR_ERROR: str | None = None

MAX_UPLOAD_BYTES = int(os.getenv("PDFSNITCH_MAX_UPLOAD_MB", "50")) * 1024 * 1024
MAX_REQUEST_BYTES = MAX_UPLOAD_BYTES * 10
TEMP_ROOT = Path(os.getenv("PDFSNITCH_TEMP_DIR", str(Path(tempfile.gettempdir()) / "pdfsnitch"))).resolve()
TEMP_ROOT.mkdir(parents=True, exist_ok=True)

PDF_EXTENSIONS = {".pdf"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp"}
DOCX_EXTENSIONS = {".docx"}

app = FastAPI(title="PDFSnitch API", version="1.0.0")
origins = [item.strip() for item in os.getenv(
    "PDFSNITCH_FRONTEND_ORIGINS",
    "http://127.0.0.1:4173,http://localhost:4173,http://127.0.0.1:5173,http://localhost:5173",
).split(",") if item.strip()]
origin_regex = os.getenv("PDFSNITCH_FRONTEND_ORIGIN_REGEX", r"https://.*\.vercel\.app").strip() or None
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_origin_regex=origin_regex,
    allow_methods=["GET", "POST", "DELETE"],
    allow_headers=["*"],
)


@app.middleware("http")
async def reject_oversized_requests(request: Request, call_next):
    content_length = request.headers.get("content-length")
    if content_length and int(content_length) > MAX_REQUEST_BYTES:
        return JSONResponse({"detail": "The upload is too large."}, status_code=413)
    return await call_next(request)


@app.exception_handler(Exception)
async def unexpected_error(_request: Request, _error: Exception):
    logger.exception("Unhandled PDFSnitch processing error", exc_info=_error)
    return JSONResponse({"detail": "Processing failed. Check the file and try again."}, status_code=500)


@contextmanager
def isolated_job():
    with tempfile.TemporaryDirectory(prefix="job-", dir=TEMP_ROOT) as directory:
        yield Path(directory)


def safe_stem(filename: str | None, fallback: str = "document") -> str:
    stem = Path(filename or fallback).stem
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", stem).strip(".-")
    return cleaned[:80] or fallback


async def read_upload(upload: UploadFile, allowed: set[str]) -> tuple[bytes, str]:
    filename = Path(upload.filename or "upload").name
    extension = Path(filename).suffix.lower()
    if extension not in allowed:
        raise HTTPException(415, f"Unsupported file type: {extension or 'unknown'}")
    data = await upload.read(MAX_UPLOAD_BYTES + 1)
    if not data:
        raise HTTPException(400, "The uploaded file is empty.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(413, f"Each file must be {MAX_UPLOAD_BYTES // 1024 // 1024} MB or smaller.")
    with isolated_job() as job:
        upload_path = job / f"{uuid.uuid4().hex}{extension}"
        upload_path.write_bytes(data)
        if extension == ".pdf":
            if not data.startswith(b"%PDF"):
                raise HTTPException(415, "The uploaded file is not a valid PDF.")
            try:
                reader = PdfReader(upload_path)
                if not reader.pages:
                    raise ValueError("No pages")
            except Exception as exc:
                raise HTTPException(422, "The PDF is damaged, empty, or password protected.") from exc
        elif extension in IMAGE_EXTENSIONS:
            try:
                with Image.open(upload_path) as image:
                    image.verify()
            except (UnidentifiedImageError, OSError) as exc:
                raise HTTPException(415, "The uploaded file is not a valid image.") from exc
        elif extension == ".docx":
            if not data.startswith(b"PK"):
                raise HTTPException(415, "The uploaded file is not a valid DOCX file.")
            try:
                document = Document(upload_path)
                if not document.paragraphs and not document.tables:
                    raise ValueError("No readable content")
            except Exception as exc:
                raise HTTPException(422, "The Word document is damaged, empty, or unsupported. Please upload a .docx file.") from exc
    return data, filename


def parse_pages(value: str, page_count: int) -> list[int]:
    if not value.strip():
        raise HTTPException(400, "Enter at least one page number or range.")
    pages: set[int] = set()
    for token in value.replace(" ", "").split(","):
        if not token:
            continue
        if "-" in token:
            pieces = token.split("-", 1)
            if not all(piece.isdigit() for piece in pieces):
                raise HTTPException(400, f"Invalid page range: {token}")
            start, end = map(int, pieces)
            if start > end:
                raise HTTPException(400, f"Page range must increase: {token}")
            pages.update(range(start, end + 1))
        elif token.isdigit():
            pages.add(int(token))
        else:
            raise HTTPException(400, f"Invalid page number: {token}")
    if not pages or min(pages) < 1 or max(pages) > page_count:
        raise HTTPException(400, f"Pages must be between 1 and {page_count}.")
    return sorted(page - 1 for page in pages)


def pdf_bytes(writer: PdfWriter) -> bytes:
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def attachment(data: bytes, filename: str, media_type: str, headers: dict[str, str] | None = None) -> Response:
    response_headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    response_headers.update(headers or {})
    return Response(data, media_type=media_type, headers=response_headers)


def zip_files(entries: Iterable[tuple[str, bytes]]) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        for filename, data in entries:
            archive.writestr(filename, data)
    return output.getvalue()


def docx_to_pdf_bytes(data: bytes) -> bytes:
    source = Document(io.BytesIO(data))
    output = io.BytesIO()
    pdf = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=42,
        leftMargin=42,
        topMargin=48,
        bottomMargin=48,
        title="PDFSnitch Word conversion",
    )
    styles = getSampleStyleSheet()
    story = []

    for paragraph in source.paragraphs:
        text = paragraph.text.strip()
        if not text:
            story.append(Spacer(1, 8))
            continue
        style_name = "Heading1" if paragraph.style and "Heading" in paragraph.style.name else "BodyText"
        story.append(Paragraph(html.escape(text), styles[style_name]))
        story.append(Spacer(1, 8))

    for table in source.tables:
        rows = []
        for row in table.rows:
            rows.append([Paragraph(html.escape(cell.text.strip() or " "), styles["BodyText"]) for cell in row.cells])
        if rows:
            story.append(Spacer(1, 10))
            table_node = Table(rows, repeatRows=1)
            table_node.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#bfd8d2")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8fff8")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.append(table_node)
            story.append(Spacer(1, 12))

    if not story:
        story.append(Paragraph("No readable text found in this Word document.", styles["BodyText"]))
    pdf.build(story)
    return output.getvalue()


def normalized_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def page_has_selectable_text(page: fitz.Page) -> bool:
    text = normalized_text(page.get_text("text"))
    return len(text) >= 12


def extract_page_blocks(page: fitz.Page, textpage=None) -> list[str]:
    blocks = page.get_text("blocks", sort=True, textpage=textpage) if textpage else page.get_text("blocks", sort=True)
    content: list[str] = []
    for block in blocks:
        if len(block) < 5:
            continue
        text = str(block[4] or "").strip()
        if not text:
            continue
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if lines:
            content.append("\n".join(lines))
    if not content:
        fallback = page.get_text("text", textpage=textpage).strip() if textpage else page.get_text("text").strip()
        if fallback:
            content.append(fallback)
    return content


def tesseract_status() -> dict[str, Any]:
    command = os.getenv("TESSERACT_CMD", "tesseract").strip() or "tesseract"
    binary = shutil.which(command)
    status: dict[str, Any] = {
        "available": False,
        "binary": binary or "",
        "version": "",
        "languages": [],
        "required_language": os.getenv("PDFSNITCH_OCR_LANG", "eng").strip() or "eng",
        "tessdata_prefix": os.getenv("TESSDATA_PREFIX", ""),
        "error": "",
    }
    if not binary:
        status["error"] = f"Tesseract command not found: {command}"
        return status
    try:
        version = subprocess.run([binary, "--version"], capture_output=True, text=True, timeout=10)
        status["version"] = (version.stdout or version.stderr).splitlines()[0].strip()
    except Exception as exc:
        status["error"] = f"Cannot run Tesseract: {exc}"
        return status
    try:
        languages = subprocess.run([binary, "--list-langs"], capture_output=True, text=True, timeout=10)
        listed = [
            line.strip()
            for line in (languages.stdout or "").splitlines()
            if line.strip() and not line.lower().startswith("list of")
        ]
        status["languages"] = listed
        required = status["required_language"]
        required_parts = [part for part in re.split(r"[+,\s]+", required) if part]
        status["available"] = all(part in listed for part in required_parts) if listed else True
        if not status["available"]:
            status["error"] = f"Required OCR language is missing: {required}"
    except Exception as exc:
        status["available"] = True
        status["error"] = f"Tesseract is installed, but languages could not be listed: {exc}"
    return status


def tesseract_ocr_page_blocks(page: fitz.Page, language: str, dpi: int) -> tuple[list[str], str | None]:
    command = os.getenv("TESSERACT_CMD", "tesseract").strip() or "tesseract"
    binary = shutil.which(command)
    if not binary:
        return [], f"Tesseract command not found: {command}"

    timeout = int(os.getenv("PDFSNITCH_OCR_TIMEOUT", "90"))
    page_dpi = max(150, min(dpi, 400))
    with tempfile.TemporaryDirectory(prefix="ocr-page-", dir=TEMP_ROOT) as directory:
        image_path = Path(directory) / f"page-{page.number + 1}.png"
        clean_image_path = Path(directory) / f"page-{page.number + 1}-clean.png"
        try:
            pixmap = page.get_pixmap(dpi=page_dpi, alpha=False, colorspace=fitz.csRGB)
            pixmap.save(image_path)
            with Image.open(image_path) as image:
                cleaned = ImageOps.grayscale(image)
                cleaned = ImageOps.autocontrast(cleaned)
                cleaned = ImageEnhance.Sharpness(cleaned).enhance(1.8)
                cleaned = ImageEnhance.Contrast(cleaned).enhance(1.4)
                cleaned.save(clean_image_path)
        except Exception as exc:
            return [], f"Could not render page for OCR: {exc}"

        psm_values = [
            item.strip()
            for item in os.getenv("PDFSNITCH_OCR_PSM", "6,3,11").split(",")
            if item.strip()
        ]
        languages = [language]
        if language != "eng":
            languages.append("eng")
        image_candidates = [clean_image_path, image_path]
        best_text = ""
        errors: list[str] = []
        for candidate_language in dict.fromkeys(languages):
            for psm in dict.fromkeys(psm_values):
                for candidate_image in image_candidates:
                    command_args = [
                        binary,
                        str(candidate_image),
                        "stdout",
                        "-l",
                        candidate_language,
                        "--psm",
                        psm,
                    ]
                    try:
                        completed = subprocess.run(command_args, capture_output=True, text=True, timeout=timeout)
                    except subprocess.TimeoutExpired:
                        errors.append(f"Tesseract timed out after {timeout} seconds")
                        continue
                    except Exception as exc:
                        errors.append(f"Tesseract failed to start: {exc}")
                        continue
                    text = completed.stdout.strip()
                    if text and len(text) > len(best_text):
                        best_text = text
                    if completed.returncode == 0 and len(normalized_text(text)) >= 8:
                        blocks = [block.strip() for block in re.split(r"\n\s*\n+", text) if block.strip()]
                        return blocks or [text], None
                    if completed.returncode != 0:
                        errors.append((completed.stderr or completed.stdout or "Tesseract OCR failed.").strip())

        text = best_text.strip()
        if not text:
            return [], errors[-1] if errors else "Tesseract OCR did not detect text on this page."
        blocks = [block.strip() for block in re.split(r"\n\s*\n+", text) if block.strip()]
        return blocks or [text], None


def ocr_page_blocks(page: fitz.Page) -> tuple[list[str], str | None]:
    language = os.getenv("PDFSNITCH_OCR_LANG", "eng").strip() or "eng"
    dpi = int(os.getenv("PDFSNITCH_OCR_DPI", "200"))
    pymupdf_error: str | None = None
    try:
        textpage = page.get_textpage_ocr(language=language, dpi=max(100, min(dpi, 300)), full=True)
        blocks = extract_page_blocks(page, textpage=textpage)
        if blocks:
            return blocks, None
        pymupdf_error = "PyMuPDF OCR did not detect text on this page."
    except Exception as exc:
        pymupdf_error = str(exc)
        logger.warning("PyMuPDF OCR failed on page %s: %s", page.number + 1, exc)

    fallback_blocks, fallback_error = tesseract_ocr_page_blocks(page, language, dpi)
    if fallback_blocks:
        return fallback_blocks, None
    combined_error = fallback_error or pymupdf_error or "OCR could not detect text."
    logger.warning("Tesseract OCR failed on page %s: %s", page.number + 1, combined_error)
    return [], combined_error


def add_blocks_to_docx(output_doc: Document, blocks: list[str], source: str) -> None:
    if not blocks:
        output_doc.add_paragraph(f"[{source} could not detect clear text on this page.]")
        return
    for block in blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        first = lines[0]
        if len(lines) == 1 and len(first) <= 90 and not first.endswith("."):
            output_doc.add_heading(first, level=3)
            continue
        paragraph = output_doc.add_paragraph()
        for index, line in enumerate(lines):
            if index:
                paragraph.add_run().add_break()
            paragraph.add_run(line)


def add_hidden_text_layer(output_doc: Document, blocks: list[str]) -> None:
    if not blocks:
        return
    paragraph = output_doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run("\n".join(blocks))
    run.font.hidden = True
    run.font.size = Pt(1)


def span_color_hex(value: int | None) -> str:
    if value is None:
        return "000000"
    return f"{int(value) & 0xFFFFFF:06X}"


def is_bold_span(span: dict[str, Any]) -> bool:
    font = str(span.get("font", "")).lower()
    return "bold" in font or "black" in font or bool(int(span.get("flags", 0)) & 16)


def is_italic_span(span: dict[str, Any]) -> bool:
    font = str(span.get("font", "")).lower()
    return "italic" in font or "oblique" in font or bool(int(span.get("flags", 0)) & 2)


def xml_text(value: str) -> str:
    return html.escape(value or "", quote=True)


def editable_pdf_lines(page: fitz.Page) -> list[dict[str, Any]]:
    lines: list[dict[str, Any]] = []
    page_dict = page.get_text("dict", sort=True)
    for block in page_dict.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans = []
            text_parts = []
            for span in line.get("spans", []):
                text = str(span.get("text", "")).replace("\x00", "").strip()
                if not text:
                    continue
                text_parts.append(text)
                spans.append({
                    "text": text,
                    "size": max(5.0, min(float(span.get("size", 10)), 72.0)),
                    "font": str(span.get("font", "Arial"))[:64] or "Arial",
                    "bold": is_bold_span(span),
                    "italic": is_italic_span(span),
                    "color": span_color_hex(span.get("color")),
                })
            if not spans:
                continue
            bbox = fitz.Rect(line.get("bbox", block.get("bbox", (0, 0, 100, 14))))
            lines.append({
                "x": max(0.0, bbox.x0),
                "y": max(0.0, bbox.y0),
                "w": max(8.0, bbox.width + 4),
                "h": max(8.0, bbox.height + 3),
                "text": " ".join(text_parts),
                "spans": spans,
            })
    return lines


def tesseract_ocr_page_lines(page: fitz.Page) -> tuple[list[dict[str, Any]], str | None]:
    command = os.getenv("TESSERACT_CMD", "tesseract").strip() or "tesseract"
    binary = shutil.which(command)
    if not binary:
        return [], f"Tesseract command not found: {command}"

    language = os.getenv("PDFSNITCH_OCR_LANG", "eng").strip() or "eng"
    timeout = int(os.getenv("PDFSNITCH_OCR_TIMEOUT", "90"))
    dpi = max(150, min(int(os.getenv("PDFSNITCH_OCR_DPI", "300")), 400))
    with tempfile.TemporaryDirectory(prefix="ocr-layout-", dir=TEMP_ROOT) as directory:
        image_path = Path(directory) / f"page-{page.number + 1}.png"
        clean_image_path = Path(directory) / f"page-{page.number + 1}-clean.png"
        try:
            pixmap = page.get_pixmap(dpi=dpi, alpha=False, colorspace=fitz.csRGB)
            pixmap.save(image_path)
            with Image.open(image_path) as image:
                cleaned = ImageOps.grayscale(image)
                cleaned = ImageOps.autocontrast(cleaned)
                cleaned = ImageEnhance.Sharpness(cleaned).enhance(1.8)
                cleaned = ImageEnhance.Contrast(cleaned).enhance(1.4)
                cleaned.save(clean_image_path)
        except Exception as exc:
            return [], f"Could not render page for OCR: {exc}"

        psm_values = [
            item.strip()
            for item in os.getenv("PDFSNITCH_OCR_PSM", "6,3,11").split(",")
            if item.strip()
        ]
        best_rows: list[dict[str, str]] = []
        errors: list[str] = []
        for psm in dict.fromkeys(psm_values):
            args = [binary, str(clean_image_path), "stdout", "-l", language, "--psm", psm, "tsv"]
            try:
                completed = subprocess.run(args, capture_output=True, text=True, timeout=timeout)
            except subprocess.TimeoutExpired:
                errors.append(f"Tesseract layout OCR timed out after {timeout} seconds")
                continue
            except Exception as exc:
                errors.append(f"Tesseract layout OCR failed to start: {exc}")
                continue
            if completed.returncode != 0:
                errors.append((completed.stderr or completed.stdout or "Tesseract OCR failed.").strip())
                continue
            rows = []
            lines = completed.stdout.splitlines()
            if not lines:
                continue
            headers = lines[0].split("\t")
            for raw in lines[1:]:
                cells = raw.split("\t")
                if len(cells) < len(headers):
                    continue
                row = dict(zip(headers, cells))
                text = row.get("text", "").strip()
                if not text:
                    continue
                try:
                    confidence = float(row.get("conf", "-1"))
                except ValueError:
                    confidence = -1
                if confidence < 0:
                    continue
                rows.append(row)
            if len(rows) > len(best_rows):
                best_rows = rows
            if rows:
                break

    if not best_rows:
        return [], errors[-1] if errors else "Tesseract OCR did not detect positioned text."

    grouped: dict[tuple[str, str, str], list[dict[str, str]]] = {}
    for row in best_rows:
        key = (row.get("block_num", "0"), row.get("par_num", "0"), row.get("line_num", "0"))
        grouped.setdefault(key, []).append(row)

    scale = 72 / dpi
    lines_out: list[dict[str, Any]] = []
    for rows in grouped.values():
        words = [row.get("text", "").strip() for row in rows if row.get("text", "").strip()]
        if not words:
            continue
        lefts = [int(float(row.get("left", "0"))) for row in rows]
        tops = [int(float(row.get("top", "0"))) for row in rows]
        rights = [int(float(row.get("left", "0"))) + int(float(row.get("width", "0"))) for row in rows]
        bottoms = [int(float(row.get("top", "0"))) + int(float(row.get("height", "0"))) for row in rows]
        x0, y0, x1, y1 = min(lefts) * scale, min(tops) * scale, max(rights) * scale, max(bottoms) * scale
        height = max(8.0, y1 - y0 + 3)
        text = " ".join(words)
        lines_out.append({
            "x": max(0.0, x0),
            "y": max(0.0, y0),
            "w": max(8.0, x1 - x0 + 4),
            "h": height,
            "text": text,
            "spans": [{
                "text": text,
                "size": max(6.0, min(height * 0.72, 18.0)),
                "font": "Arial",
                "bold": False,
                "italic": False,
                "color": "000000",
            }],
        })
    return sorted(lines_out, key=lambda item: (item["y"], item["x"])), None


def paddle_ocr_status() -> dict[str, Any]:
    if os.getenv("PDFSNITCH_PADDLEOCR_ENABLED", "1").lower() in {"0", "false", "no"}:
        return {"available": False, "error": "PaddleOCR disabled by PDFSNITCH_PADDLEOCR_ENABLED"}
    if _PADDLE_OCR_ENGINE is not None:
        return {
            "available": True,
            "loaded": True,
            "paddle_version": safe_package_version("paddlepaddle"),
            "paddleocr_version": safe_package_version("paddleocr"),
            "lang": os.getenv("PDFSNITCH_PADDLEOCR_LANG", "en"),
        }
    if importlib_util.find_spec("paddleocr") is None:
        return {"available": False, "error": "paddleocr package is not installed"}
    if importlib_util.find_spec("paddle") is None:
        return {"available": False, "error": "paddlepaddle package is not installed"}
    if _PADDLE_OCR_ERROR:
        return {
            "available": False,
            "error": _PADDLE_OCR_ERROR,
            "paddle_version": safe_package_version("paddlepaddle"),
            "paddleocr_version": safe_package_version("paddleocr"),
            "lang": os.getenv("PDFSNITCH_PADDLEOCR_LANG", "en"),
        }
    return {
        "available": True,
        "loaded": False,
        "paddle_version": safe_package_version("paddlepaddle"),
        "paddleocr_version": safe_package_version("paddleocr"),
        "lang": os.getenv("PDFSNITCH_PADDLEOCR_LANG", "en"),
    }


def safe_package_version(package_name: str) -> str:
    try:
        return importlib_metadata.version(package_name)
    except Exception:
        return ""


def get_paddle_ocr_engine():
    global _PADDLE_OCR_ENGINE, _PADDLE_OCR_ERROR
    if _PADDLE_OCR_ENGINE is not None:
        return _PADDLE_OCR_ENGINE
    if os.getenv("PDFSNITCH_PADDLEOCR_ENABLED", "1").lower() in {"0", "false", "no"}:
        _PADDLE_OCR_ERROR = "PaddleOCR disabled by PDFSNITCH_PADDLEOCR_ENABLED"
        return None
    try:
        from paddleocr import PaddleOCR  # type: ignore
        lang = os.getenv("PDFSNITCH_PADDLEOCR_LANG", "en").strip() or "en"
        try:
            _PADDLE_OCR_ENGINE = PaddleOCR(
                lang=lang,
                use_doc_orientation_classify=False,
                use_doc_unwarping=False,
                use_textline_orientation=True,
                text_det_limit_side_len=int(os.getenv("PDFSNITCH_PADDLEOCR_LIMIT", "1280")),
            )
        except TypeError:
            _PADDLE_OCR_ENGINE = PaddleOCR(lang=lang, use_angle_cls=True, show_log=False)
        _PADDLE_OCR_ERROR = None
        return _PADDLE_OCR_ENGINE
    except Exception as exc:
        _PADDLE_OCR_ERROR = str(exc)
        logger.warning("PaddleOCR unavailable: %s", exc)
        return None


def polygon_bounds(points: Any) -> tuple[float, float, float, float] | None:
    try:
        coords = []
        for point in points:
            if isinstance(point, (list, tuple)) and len(point) >= 2:
                coords.append((float(point[0]), float(point[1])))
        if not coords:
            return None
        xs = [item[0] for item in coords]
        ys = [item[1] for item in coords]
        return min(xs), min(ys), max(xs), max(ys)
    except Exception:
        return None


def normalize_paddle_items(raw_result: Any) -> list[tuple[Any, str, float]]:
    items: list[tuple[Any, str, float]] = []

    def consume_dict(data: dict[str, Any]) -> None:
        texts = data.get("rec_texts") or data.get("texts") or data.get("text") or []
        scores = data.get("rec_scores") or data.get("scores") or []
        boxes = data.get("rec_polys") or data.get("rec_boxes") or data.get("dt_polys") or data.get("boxes") or []
        if isinstance(texts, str):
            texts = [texts]
        for index, text in enumerate(texts):
            if not str(text).strip():
                continue
            box = boxes[index] if index < len(boxes) else None
            score = float(scores[index]) if index < len(scores) else 0.85
            items.append((box, str(text), score))

    def walk(value: Any) -> None:
        if value is None:
            return
        if hasattr(value, "res") and isinstance(value.res, dict):
            consume_dict(value.res)
            return
        if isinstance(value, dict):
            consume_dict(value)
            return
        if isinstance(value, (list, tuple)):
            if len(value) >= 2 and isinstance(value[1], (list, tuple)) and len(value[1]) >= 2 and isinstance(value[1][0], str):
                box, payload = value[0], value[1]
                try:
                    score = float(payload[1])
                except Exception:
                    score = 0.85
                items.append((box, str(payload[0]), score))
                return
            for child in value:
                walk(child)

    walk(raw_result)
    return items


def lines_from_paddle_image(image_path: Path, page_width: float, page_height: float) -> tuple[list[dict[str, Any]], str | None, float]:
    engine = get_paddle_ocr_engine()
    if engine is None:
        return [], _PADDLE_OCR_ERROR or "PaddleOCR unavailable", 0.0
    try:
        if hasattr(engine, "predict"):
            raw = engine.predict(str(image_path))
        else:
            raw = engine.ocr(str(image_path), cls=True)
    except Exception as exc:
        return [], str(exc), 0.0

    try:
        with Image.open(image_path) as image:
            image_width, image_height = image.size
    except Exception:
        image_width, image_height = 1, 1
    scale_x = page_width / max(1, image_width)
    scale_y = page_height / max(1, image_height)
    lines: list[dict[str, Any]] = []
    scores: list[float] = []
    for box, text, score in normalize_paddle_items(raw):
        bounds = polygon_bounds(box)
        if bounds:
            x0, y0, x1, y1 = bounds
            x, y, width, height = x0 * scale_x, y0 * scale_y, max(8, (x1 - x0) * scale_x), max(8, (y1 - y0) * scale_y)
        else:
            x, y, width, height = 0.0, len(lines) * 14.0, page_width * 0.9, 14.0
        font_size = max(6.0, min(height * 0.72, 22.0))
        clean_text = normalized_text(text)
        if not clean_text:
            continue
        scores.append(max(0.0, min(float(score), 1.0)))
        lines.append({
            "x": x,
            "y": y,
            "w": width + 4,
            "h": height + 3,
            "text": clean_text,
            "score": score,
            "spans": [{
                "text": clean_text,
                "size": font_size,
                "font": "Arial",
                "bold": False,
                "italic": False,
                "color": "000000",
            }],
        })
    confidence = sum(scores) / len(scores) if scores else 0.0
    return sorted(lines, key=lambda item: (item["y"], item["x"])), None if lines else "PaddleOCR did not detect text.", confidence


def render_page_image(page: fitz.Page, dpi: int | None = None) -> Path:
    target_dpi = dpi or max(150, min(int(os.getenv("PDFSNITCH_OCR_DPI", "300")), 400))
    directory = Path(tempfile.mkdtemp(prefix="ocr-page-", dir=TEMP_ROOT))
    image_path = directory / f"page-{page.number + 1}.png"
    pixmap = page.get_pixmap(dpi=target_dpi, alpha=False, colorspace=fitz.csRGB)
    pixmap.save(image_path)
    return image_path


def ocr_page_lines(page: fitz.Page, metadata: dict[str, Any]) -> tuple[list[dict[str, Any]], str | None, float]:
    image_path = render_page_image(page)
    try:
        paddle_lines, paddle_error, paddle_confidence = lines_from_paddle_image(image_path, page.rect.width, page.rect.height)
        if paddle_lines:
            metadata["ocr_engine"] = "paddleocr"
            return paddle_lines, None, paddle_confidence
        fallback_lines, fallback_error = tesseract_ocr_page_lines(page)
        if fallback_lines:
            metadata["ocr_engine"] = "tesseract"
            return fallback_lines, None, 0.62
        return [], paddle_error or fallback_error or "OCR did not detect text.", 0.0
    finally:
        try:
            shutil.rmtree(image_path.parent, ignore_errors=True)
        except Exception:
            pass


def textbox_run_xml(span: dict[str, Any]) -> str:
    bold = "<w:b/>" if span.get("bold") else ""
    italic = "<w:i/>" if span.get("italic") else ""
    size_half_points = int(round(float(span.get("size", 10)) * 2))
    font_name = xml_text(str(span.get("font", "Arial")) or "Arial")
    text = xml_text(str(span.get("text", "")))
    return (
        f"<w:r><w:rPr><w:rFonts w:ascii=\"{font_name}\" w:hAnsi=\"{font_name}\"/>"
        f"<w:sz w:val=\"{size_half_points}\"/><w:color w:val=\"{span.get('color', '000000')}\"/>"
        f"{bold}{italic}</w:rPr><w:t xml:space=\"preserve\">{text}</w:t></w:r>"
    )


def add_editable_textbox(paragraph, shape_id: int, line: dict[str, Any]) -> None:
    x = float(line.get("x", 0))
    y = float(line.get("y", 0))
    width = max(8.0, float(line.get("w", 120)))
    height = max(8.0, float(line.get("h", 14)))
    runs = "".join(textbox_run_xml(span) for span in line.get("spans", []))
    if not runs:
        return
    xml = (
        f"<w:r {nsdecls('w')} xmlns:v=\"urn:schemas-microsoft-com:vml\"><w:pict>"
        f"<v:shape id=\"pdfsnitch_text_{shape_id}\" type=\"#_x0000_t202\" "
        f"style=\"position:absolute;margin-left:{x:.2f}pt;margin-top:{y:.2f}pt;"
        f"width:{width:.2f}pt;height:{height:.2f}pt;z-index:{shape_id};"
        f"mso-position-horizontal:absolute;mso-position-horizontal-relative:page;"
        f"mso-position-vertical:absolute;mso-position-vertical-relative:page\" "
        f"filled=\"f\" stroked=\"f\">"
        f"<v:textbox inset=\"0,0,0,0\"><w:txbxContent>"
        f"<w:p><w:pPr><w:spacing w:before=\"0\" w:after=\"0\" w:line=\"240\" w:lineRule=\"auto\"/></w:pPr>{runs}</w:p>"
        f"</w:txbxContent></v:textbox></v:shape></w:pict></w:r>"
    )
    paragraph._p.append(parse_xml(xml))


def add_editable_page(output_doc: Document, page: fitz.Page, lines: list[dict[str, Any]], shape_start: int) -> int:
    paragraph = output_doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    for offset, line in enumerate(lines):
        add_editable_textbox(paragraph, shape_start + offset, line)
    return shape_start + len(lines)


def average_font_size(lines: list[dict[str, Any]]) -> float:
    sizes = [float(span.get("size", 10)) for line in lines for span in line.get("spans", [])]
    return sum(sizes) / len(sizes) if sizes else 10.0


def add_line_runs(paragraph, line: dict[str, Any]) -> None:
    for span in line.get("spans", []) or [{"text": line.get("text", ""), "size": 10, "font": "Arial"}]:
        run = paragraph.add_run(str(span.get("text", "")))
        run.font.name = str(span.get("font", "Arial") or "Arial")
        run.font.size = Pt(max(6, min(float(span.get("size", 10)), 48)))
        run.bold = bool(span.get("bold"))
        run.italic = bool(span.get("italic"))
        color = str(span.get("color", "000000"))
        if re.match(r"^[0-9A-Fa-f]{6}$", color):
            run.font.color.rgb = RGBColor.from_string(color.upper())


def paragraph_alignment(line: dict[str, Any], page_width: float) -> Any:
    center = float(line.get("x", 0)) + float(line.get("w", 0)) / 2
    if abs(center - page_width / 2) < page_width * 0.08 and float(line.get("w", 0)) < page_width * 0.75:
        return WD_ALIGN_PARAGRAPH.CENTER
    if float(line.get("x", 0)) > page_width * 0.58:
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def group_lines_into_blocks(lines: list[dict[str, Any]]) -> list[list[dict[str, Any]]]:
    if not lines:
        return []
    sorted_lines = sorted(lines, key=lambda item: (float(item.get("y", 0)), float(item.get("x", 0))))
    blocks: list[list[dict[str, Any]]] = []
    current: list[dict[str, Any]] = []
    previous_bottom = None
    previous_x = 0.0
    for line in sorted_lines:
        y = float(line.get("y", 0))
        h = float(line.get("h", 12))
        x = float(line.get("x", 0))
        gap = 0 if previous_bottom is None else y - previous_bottom
        new_block = previous_bottom is not None and (gap > max(10, h * 0.9) or abs(x - previous_x) > 180)
        if new_block and current:
            blocks.append(current)
            current = []
        current.append(line)
        previous_bottom = y + h
        previous_x = x
    if current:
        blocks.append(current)
    return blocks


def line_is_list_item(text: str) -> bool:
    return bool(re.match(r"^\s*(?:[-•*]|\(?[0-9]{1,2}[\).]|[A-Za-z][\).])\s+", text))


def possible_table_rows(lines: list[dict[str, Any]]) -> list[list[dict[str, Any]]]:
    rows: list[list[dict[str, Any]]] = []
    for line in sorted(lines, key=lambda item: (float(item.get("y", 0)), float(item.get("x", 0)))):
        placed = False
        y = float(line.get("y", 0))
        for row in rows:
            row_y = sum(float(item.get("y", 0)) for item in row) / len(row)
            if abs(y - row_y) <= max(5.5, float(line.get("h", 10)) * 0.45):
                row.append(line)
                placed = True
                break
        if not placed:
            rows.append([line])
    rows = [sorted(row, key=lambda item: float(item.get("x", 0))) for row in rows]
    return [row for row in rows if len(row) >= 2]


def looks_like_table(lines: list[dict[str, Any]]) -> bool:
    rows = possible_table_rows(lines)
    if len(rows) < 2:
        return False
    multi_rows = sum(1 for row in rows if len(row) >= 2)
    return multi_rows >= 2 and max(len(row) for row in rows) <= 8


def add_table_from_lines(output_doc: Document, lines: list[dict[str, Any]]) -> None:
    rows = possible_table_rows(lines)
    columns = max(len(row) for row in rows)
    table = output_doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(columns):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            if col_index >= len(row):
                continue
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            add_line_runs(paragraph, row[col_index])


def add_structured_page(output_doc: Document, lines: list[dict[str, Any]], page_width: float, page_height: float) -> None:
    top_margin = page_height * 0.08
    bottom_margin = page_height * 0.92
    header_lines = [line for line in lines if float(line.get("y", 0)) <= top_margin]
    footer_lines = [line for line in lines if float(line.get("y", 0)) >= bottom_margin]
    body_lines = [line for line in lines if line not in header_lines and line not in footer_lines]
    ordered_groups = [header_lines] + group_lines_into_blocks(body_lines) + [footer_lines]
    for block in [group for group in ordered_groups if group]:
        if looks_like_table(block):
            add_table_from_lines(output_doc, block)
            output_doc.add_paragraph()
            continue
        for line in block:
            text = str(line.get("text", "")).strip()
            if not text:
                continue
            style = None
            if average_font_size([line]) >= 17 or (line.get("spans") and line["spans"][0].get("bold") and len(text) <= 90):
                style = "Heading 2"
            paragraph = output_doc.add_paragraph(style=style)
            paragraph.alignment = paragraph_alignment(line, page_width)
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(3)
            if line_is_list_item(text):
                paragraph.style = "List Bullet"
            add_line_runs(paragraph, line)


def layout_confidence(metadata: dict[str, Any], all_lines: list[list[dict[str, Any]]], high_fidelity: bool) -> float:
    page_count = max(1, int(metadata.get("page_count", 1)))
    failed = len(metadata.get("ocr_failed_pages", [])) + len(metadata.get("image_fallback_pages", []))
    text_pages = len(metadata.get("text_pages", []))
    line_count = sum(len(page_lines) for page_lines in all_lines)
    base = 0.68
    if text_pages:
        base += 0.18 * (text_pages / page_count)
    if metadata.get("ocr_engine") == "paddleocr":
        base += 0.1
    if high_fidelity:
        base += 0.04
    if line_count < page_count * 3:
        base -= 0.2
    base -= 0.18 * (failed / page_count)
    return max(0.05, min(base, 0.98))


def collect_page_text_blocks(page: fitz.Page, metadata: dict[str, Any], page_index: int) -> list[str]:
    if page_has_selectable_text(page):
        metadata["text_pages"].append(page_index)
        return extract_page_blocks(page)
    metadata["ocr_pages"].append(page_index)
    blocks, ocr_error = ocr_page_blocks(page)
    if ocr_error:
        metadata["ocr_failed_pages"].append(page_index)
    return blocks


def image_to_pdf_document(data: bytes) -> fitz.Document:
    image = Image.open(io.BytesIO(data))
    image.load()
    width_px, height_px = image.size
    page_width = 612.0
    page_height = max(180.0, page_width * height_px / max(1, width_px))
    image_stream = io.BytesIO()
    image.convert("RGB").save(image_stream, format="PNG")
    document = fitz.open()
    page = document.new_page(width=page_width, height=page_height)
    page.insert_image(page.rect, stream=image_stream.getvalue())
    return document


def source_to_fitz_document(data: bytes, extension: str) -> fitz.Document:
    if extension == ".pdf":
        return fitz.open(stream=data, filetype="pdf")
    if extension in IMAGE_EXTENSIONS:
        return image_to_pdf_document(data)
    raise HTTPException(415, f"Unsupported file type: {extension or 'unknown'}")


def pdf_to_docx_bytes(data: bytes, extension: str = ".pdf", high_fidelity: bool = False) -> tuple[bytes, dict[str, Any]]:
    document = source_to_fitz_document(data, extension)
    metadata = {
        "text_pages": [],
        "ocr_pages": [],
        "ocr_failed_pages": [],
        "image_fallback_pages": [],
        "page_count": 0,
        "layout_mode": "high_fidelity_editable_textboxes" if high_fidelity else "editable_document_reconstruction",
        "ocr_engine": "none",
        "high_fidelity": high_fidelity,
    }
    all_page_lines: list[list[dict[str, Any]]] = []
    try:
        output_doc = Document()
        if document.page_count:
            first_page = document[0]
            first_rect = first_page.rect
            section = output_doc.sections[0]
            section.page_width = Inches(first_rect.width / 72)
            section.page_height = Inches(first_rect.height / 72)
            margin = 0 if high_fidelity else 0.45
            section.top_margin = Inches(margin)
            section.bottom_margin = Inches(margin)
            section.left_margin = Inches(margin)
            section.right_margin = Inches(margin)
            section.header_distance = Inches(0)
            section.footer_distance = Inches(0)
        styles = output_doc.styles
        normal = styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(3)
        normal.paragraph_format.line_spacing = 1
        metadata["page_count"] = document.page_count
        shape_id = 1
        for page_index, page in enumerate(document, start=1):
            if page_index > 1:
                output_doc.add_page_break()
            if page_has_selectable_text(page):
                metadata["text_pages"].append(page_index)
                lines = editable_pdf_lines(page)
                confidence = 0.94
            else:
                metadata["ocr_pages"].append(page_index)
                lines, ocr_error, confidence = ocr_page_lines(page, metadata)
                if ocr_error:
                    metadata["ocr_failed_pages"].append(page_index)
            if not lines:
                metadata["image_fallback_pages"].append(page_index)
                note = output_doc.add_paragraph("[PDFSnitch could not detect editable text on this page.]")
                note.runs[0].font.size = Pt(10)
                all_page_lines.append([])
                continue
            all_page_lines.append(lines)
            if high_fidelity:
                shape_id = add_editable_page(output_doc, page, lines, shape_id)
            else:
                add_structured_page(output_doc, lines, page.rect.width, page.rect.height)
            metadata.setdefault("page_confidences", []).append(round(confidence, 3))
    finally:
        document.close()
    metadata["layout_confidence"] = round(layout_confidence(metadata, all_page_lines, high_fidelity), 3)
    output = io.BytesIO()
    output_doc.save(output)
    return output.getvalue(), metadata


@app.get("/api/health")
def health():
    return {"status": "ok", "max_upload_mb": MAX_UPLOAD_BYTES // 1024 // 1024}


@app.get("/api/ocr-status")
def ocr_status():
    return {"paddleocr": paddle_ocr_status(), "tesseract": tesseract_status()}


@app.post("/api/preview")
async def preview_pdf(file: UploadFile = File(...), max_pages: int = Form(50)):
    data, _ = await read_upload(file, PDF_EXTENSIONS)
    try:
        document = fitz.open(stream=data, filetype="pdf")
        limit = min(max(1, max_pages), 100, document.page_count)
        previews = []
        for index in range(limit):
            pixmap = document[index].get_pixmap(matrix=fitz.Matrix(0.8, 0.8), alpha=False)
            encoded = base64.b64encode(pixmap.tobytes("png")).decode("ascii")
            previews.append({"page": index + 1, "src": f"data:image/png;base64,{encoded}"})
        return {"page_count": document.page_count, "previews": previews, "truncated": limit < document.page_count}
    except Exception as exc:
        raise HTTPException(422, "Unable to render this PDF.") from exc


@app.post("/api/split")
async def split_pdf(file: UploadFile = File(...), mode: str = Form("individual"), ranges: str = Form("")):
    data, filename = await read_upload(file, PDF_EXTENSIONS)
    reader = PdfReader(io.BytesIO(data))
    stem = safe_stem(filename)
    if mode == "individual":
        outputs = []
        for index, page in enumerate(reader.pages, start=1):
            writer = PdfWriter()
            writer.add_page(page)
            outputs.append((f"{stem}-page-{index}.pdf", pdf_bytes(writer)))
        return attachment(zip_files(outputs), f"{stem}-split-pages.zip", "application/zip")
    if mode != "ranges":
        raise HTTPException(400, "Split mode must be individual or ranges.")
    selected = parse_pages(ranges, len(reader.pages))
    writer = PdfWriter()
    for page_index in selected:
        writer.add_page(reader.pages[page_index])
    return attachment(pdf_bytes(writer), f"{stem}-pages.pdf", "application/pdf")


@app.post("/api/pdf-to-images")
async def pdf_to_images(file: UploadFile = File(...), format: str = Form("png"), dpi: int = Form(150)):
    data, filename = await read_upload(file, PDF_EXTENSIONS)
    image_format = format.lower()
    if image_format not in {"png", "jpg", "jpeg"}:
        raise HTTPException(400, "Image format must be PNG or JPG.")
    dpi = min(max(dpi, 72), 300)
    stem = safe_stem(filename)
    entries = []
    extension = "jpg" if image_format in {"jpg", "jpeg"} else "png"
    try:
        document = fitz.open(stream=data, filetype="pdf")
        try:
            for index, page in enumerate(document, start=1):
                last_error = None
                for render_dpi in dict.fromkeys((dpi, min(dpi, 150), 96)):
                    try:
                        pixmap = page.get_pixmap(dpi=render_dpi, alpha=False, colorspace=fitz.csRGB)
                        payload = pixmap.tobytes("jpeg", jpg_quality=90) if extension == "jpg" else pixmap.tobytes("png")
                        entries.append((f"{stem}-page-{index}.{extension}", payload))
                        break
                    except Exception as exc:
                        last_error = exc
                else:
                    raise HTTPException(422, f"Page {index} could not be converted. Try a lower resolution.") from last_error
        finally:
            document.close()
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("PDF-to-image conversion failed for %s", filename)
        raise HTTPException(422, "This PDF could not be converted. It may contain unsupported or damaged page data.") from exc
    return attachment(zip_files(entries), f"{stem}-{extension}-images.zip", "application/zip")


@app.post("/api/images-to-pdf")
async def images_to_pdf(files: list[UploadFile] = File(...)):
    if not files:
        raise HTTPException(400, "Choose at least one image.")
    images: list[Image.Image] = []
    try:
        for upload in files:
            data, _ = await read_upload(upload, IMAGE_EXTENSIONS)
            with Image.open(io.BytesIO(data)) as source:
                frame = source.convert("RGB")
                frame.load()
                images.append(frame)
        output = io.BytesIO()
        images[0].save(output, "PDF", save_all=True, append_images=images[1:], resolution=150, quality=95)
        return attachment(output.getvalue(), f"images-{uuid.uuid4().hex[:8]}.pdf", "application/pdf")
    finally:
        for image in images:
            image.close()


@app.post("/api/word-to-pdf")
async def word_to_pdf(file: UploadFile = File(...)):
    data, filename = await read_upload(file, DOCX_EXTENSIONS)
    converted = docx_to_pdf_bytes(data)
    return attachment(converted, f"{safe_stem(filename)}.pdf", "application/pdf")


@app.post("/api/pdf-to-word")
async def pdf_to_word(file: UploadFile = File(...), high_fidelity: bool = Form(False)):
    allowed_extensions = PDF_EXTENSIONS | IMAGE_EXTENSIONS
    data, filename = await read_upload(file, allowed_extensions)
    extension = Path(filename).suffix.lower()
    converted, metadata = await run_in_threadpool(pdf_to_docx_bytes, data, extension, high_fidelity)
    headers = {
        "X-PDFSNITCH-Text-Pages": ",".join(map(str, metadata["text_pages"])),
        "X-PDFSNITCH-OCR-Pages": ",".join(map(str, metadata["ocr_pages"])),
        "X-PDFSNITCH-OCR-Failed-Pages": ",".join(map(str, metadata["ocr_failed_pages"])),
        "X-PDFSNITCH-Image-Fallback-Pages": ",".join(map(str, metadata["image_fallback_pages"])),
        "X-PDFSNITCH-Layout-Mode": str(metadata.get("layout_mode", "editable_document_reconstruction")),
        "X-PDFSNITCH-Layout-Confidence": str(metadata.get("layout_confidence", 0)),
        "X-PDFSNITCH-OCR-Engine": str(metadata.get("ocr_engine", "none")),
        "X-PDFSNITCH-High-Fidelity": "1" if metadata.get("high_fidelity") else "0",
        "Access-Control-Expose-Headers": "Content-Disposition, X-PDFSNITCH-Text-Pages, X-PDFSNITCH-OCR-Pages, X-PDFSNITCH-OCR-Failed-Pages, X-PDFSNITCH-Image-Fallback-Pages, X-PDFSNITCH-Layout-Mode, X-PDFSNITCH-Layout-Confidence, X-PDFSNITCH-OCR-Engine, X-PDFSNITCH-High-Fidelity",
    }
    return attachment(
        converted,
        f"{safe_stem(filename)}.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers,
    )


@app.post("/api/compress")
async def compress_pdf(
    file: UploadFile = File(...),
    level: str = Form("medium"),
    resolution: int = Form(144),
    conversion: str = Form("None"),
    multimedia: str = Form("Discard"),
    fonts: str = Form("Leave unchanged"),
):
    data, filename = await read_upload(file, PDF_EXTENSIONS)
    presets = {"low": (150, 86), "medium": (110, 70), "high": (84, 52)}
    if level not in presets:
        raise HTTPException(400, "Compression level must be low, medium, or high.")
    preset_dpi, quality = presets[level]
    dpi = min(max(resolution or preset_dpi, 72), 300)
    grayscale = conversion.lower() == "grayscale"
    document = fitz.open(stream=data, filetype="pdf")
    try:
        document.rewrite_images(dpi_threshold=dpi + 1, dpi_target=dpi, quality=quality, lossy=True, lossless=True, set_to_gray=grayscale)
        compressed = document.tobytes(garbage=4, deflate=True, deflate_images=True, deflate_fonts=True, use_objstms=1)
    finally:
        document.close()
    if len(compressed) >= len(data):
        compressed = data
    original_size = len(data)
    compressed_size = len(compressed)
    reduction = max(0.0, (1 - compressed_size / original_size) * 100)
    headers = {
        "X-Original-Size": str(original_size),
        "X-Compressed-Size": str(compressed_size),
        "X-Compression-Percent": f"{reduction:.1f}",
        "Access-Control-Expose-Headers": "Content-Disposition, X-Original-Size, X-Compressed-Size, X-Compression-Percent",
    }
    return attachment(compressed, f"{safe_stem(filename)}-{level}-compressed.pdf", "application/pdf", headers)


@app.post("/api/delete-pages")
async def delete_pages(file: UploadFile = File(...), pages: str = Form(...)):
    data, filename = await read_upload(file, PDF_EXTENSIONS)
    reader = PdfReader(io.BytesIO(data))
    deleted = set(parse_pages(pages, len(reader.pages)))
    if len(deleted) >= len(reader.pages):
        raise HTTPException(400, "At least one page must remain in the PDF.")
    writer = PdfWriter()
    for index, page in enumerate(reader.pages):
        if index not in deleted:
            writer.add_page(page)
    return attachment(pdf_bytes(writer), f"{safe_stem(filename)}-pages-removed.pdf", "application/pdf")


@app.post("/api/merge")
async def merge_pdfs(files: list[UploadFile] = File(...)):
    if len(files) < 2:
        raise HTTPException(400, "Choose at least two PDFs to merge.")
    writer = PdfWriter()
    for upload in files:
        data, _ = await read_upload(upload, PDF_EXTENSIONS)
        writer.append(io.BytesIO(data))
    return attachment(pdf_bytes(writer), f"merged-{uuid.uuid4().hex[:8]}.pdf", "application/pdf")


@app.post("/api/rotate")
async def rotate_pdf(file: UploadFile = File(...), degrees: int = Form(90)):
    if degrees not in {-90, 90, 180}:
        raise HTTPException(400, "Rotation must be -90, 90, or 180 degrees.")
    data, filename = await read_upload(file, PDF_EXTENSIONS)
    reader = PdfReader(io.BytesIO(data))
    writer = PdfWriter()
    for page in reader.pages:
        page.rotate(degrees)
        writer.add_page(page)
    return attachment(pdf_bytes(writer), f"{safe_stem(filename)}-rotated.pdf", "application/pdf")


def watermark_page(width: float, height: float, text: str, opacity: float) -> bytes:
    output = io.BytesIO()
    layer = canvas.Canvas(output, pagesize=(width, height))
    layer.saveState()
    layer.setFillColor(Color(0.08, 0.42, 0.34, alpha=opacity))
    layer.setFont("Helvetica-Bold", max(18, min(width, height) / 13))
    layer.translate(width / 2, height / 2)
    layer.rotate(35)
    layer.drawCentredString(0, 0, text[:120])
    layer.restoreState()
    layer.save()
    return output.getvalue()


@app.post("/api/watermark")
async def watermark_pdf(file: UploadFile = File(...), text: str = Form(...), opacity: int = Form(35)):
    if not text.strip():
        raise HTTPException(400, "Enter watermark text.")
    data, filename = await read_upload(file, PDF_EXTENSIONS)
    reader = PdfReader(io.BytesIO(data))
    writer = PdfWriter()
    alpha = min(max(opacity, 10), 100) / 100
    for page in reader.pages:
        box = page.mediabox
        overlay = PdfReader(io.BytesIO(watermark_page(float(box.width), float(box.height), text.strip(), alpha))).pages[0]
        page.merge_page(overlay)
        writer.add_page(page)
    return attachment(pdf_bytes(writer), f"{safe_stem(filename)}-watermarked.pdf", "application/pdf")


@app.post("/api/protect")
async def protect_pdf(file: UploadFile = File(...), password: str = Form(...)):
    if len(password) < 6:
        raise HTTPException(400, "Password must contain at least 6 characters.")
    data, filename = await read_upload(file, PDF_EXTENSIONS)
    reader = PdfReader(io.BytesIO(data))
    writer = PdfWriter()
    writer.append_pages_from_reader(reader)
    writer.encrypt(password)
    return attachment(pdf_bytes(writer), f"{safe_stem(filename)}-protected.pdf", "application/pdf")


@app.post("/api/unlock")
async def unlock_pdf(file: UploadFile = File(...), password: str = Form(...)):
    data = await file.read(MAX_UPLOAD_BYTES + 1)
    if len(data) > MAX_UPLOAD_BYTES:
        raise HTTPException(413, "The PDF is too large.")
    try:
        reader = PdfReader(io.BytesIO(data))
        if not reader.is_encrypted or reader.decrypt(password) == 0:
            raise HTTPException(400, "The password is incorrect or the PDF is not encrypted.")
        writer = PdfWriter()
        writer.append_pages_from_reader(reader)
        return attachment(pdf_bytes(writer), f"{safe_stem(file.filename)}-unlocked.pdf", "application/pdf")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(422, "Unable to unlock this PDF.") from exc


try:
    from .admin_api import init_admin_storage, router as admin_router
except ImportError:
    from admin_api import init_admin_storage, router as admin_router

init_admin_storage()
app.include_router(admin_router)
